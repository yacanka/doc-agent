"""Safe DOCX operations."""
from __future__ import annotations

import re
import shutil
import zipfile
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from lxml import etree

_SAFE_NAME = re.compile(r"[^A-Za-z0-9_.-]+")
_WORD_NAMESPACE = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_TEXT_TAG = f"{{{_WORD_NAMESPACE}}}t"
_PARAGRAPH_TAG = f"{{{_WORD_NAMESPACE}}}p"
_XML_PARTS = re.compile(r"word/(document|header\d+|footer\d+|footnotes|endnotes|comments)\.xml$")


@dataclass(frozen=True)
class ToolResult:
    """Result returned by a document tool operation."""
    output_path: Path
    changed_count: int
    message: str = ""


def safe_filename(filename: str) -> str:
    """Return a path-safe filename for local storage."""
    name = Path(filename).name.strip() or "document.docx"
    return _SAFE_NAME.sub("_", name)


def assert_docx(filename: str, content_type: str | None) -> None:
    """Validate that an upload is a DOCX file."""
    allowed = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    valid_type = content_type in {None, allowed, "application/octet-stream"}
    if not filename.lower().endswith(".docx") or not valid_type:
        raise ValueError("Only DOCX uploads are supported")


def extract_text(path: Path) -> str:
    """Extract visible text from paragraphs, tables, headers, and footers."""
    document = Document(path)
    parts = _paragraph_text(document.paragraphs)
    parts.extend(_table_text(document.tables))
    for section in document.sections:
        parts.extend(_paragraph_text(section.header.paragraphs))
        parts.extend(_paragraph_text(section.footer.paragraphs))
        parts.extend(_table_text(section.header.tables))
        parts.extend(_table_text(section.footer.tables))
    return "\n".join(part for part in parts if part)


def replace_text_preserve_style(source_path: Path, target: str, replacement: str, output_path: Path) -> ToolResult:
    """Replace DOCX text in a timestamped copy while preserving styles."""
    if not target:
        raise ValueError("Replacement target must not be empty")
    destination = _timestamped_copy_path(output_path)
    destination.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(source_path, destination)
    changed_count = _replace_xml_text(destination, target, replacement, split_only=True)
    if not changed_count:
        changed_count = _replace_xml_text(destination, target, replacement)
    return ToolResult(destination, changed_count, f"Replaced {changed_count} occurrence(s)")


def replace_text(source: Path, destination: Path, replacements: dict[str, str]) -> int:
    """Copy a DOCX and replace text inside paragraphs, tables, and XML parts."""
    shutil.copy2(source, destination)
    count = 0
    for target, replacement in replacements.items():
        count += _replace_xml_text(destination, target, replacement)
    return count


def timestamped_output(outputs_dir: Path, session_id: str, filename: str) -> Path:
    """Build a unique timestamped output path for a session."""
    stamp = datetime.now(timezone.utc).strftime("%Y%m%dT%H%M%SZ")
    return outputs_dir / session_id / f"{stamp}_{safe_filename(filename)}"

def _paragraph_text(paragraphs: list) -> list[str]:
    return [paragraph.text for paragraph in paragraphs]


def _table_text(tables: list) -> list[str]:
    return [cell.text for table in tables for row in table.rows for cell in row.cells]


def _timestamped_copy_path(output_path: Path) -> Path:
    stamp = datetime.now(timezone.utc).strftime("%Y%m%dT%H%M%S%fZ")
    name = safe_filename(output_path.name or "document.docx")
    return output_path.parent / f"{stamp}_{name}"


def _replace_xml_text(path: Path, target: str, replacement: str, split_only: bool = False) -> int:
    parts = _read_zip(path)
    changed, total = _replace_xml_parts(parts, target, replacement, split_only)
    if changed:
        _write_zip(path, parts)
    return total


def _read_zip(path: Path) -> dict[str, bytes]:
    with zipfile.ZipFile(path, "r") as archive:
        return {name: archive.read(name) for name in archive.namelist()}


def _write_zip(path: Path, parts: dict[str, bytes]) -> None:
    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as archive:
        for name, content in parts.items():
            archive.writestr(name, content)


def _replace_xml_parts(
    parts: dict[str, bytes], target: str, replacement: str, split_only: bool
) -> tuple[bool, int]:
    total = 0
    changed = False
    for name, content in list(parts.items()):
        if not _XML_PARTS.match(name):
            continue
        new_content, count = _replace_xml_part(content, target, replacement, split_only)
        parts[name] = new_content
        total += count
        changed = changed or bool(count)
    return changed, total


def _replace_xml_part(content: bytes, target: str, replacement: str, split_only: bool) -> tuple[bytes, int]:
    parser = etree.XMLParser(resolve_entities=False, remove_blank_text=False)
    root = etree.fromstring(content, parser=parser)
    count = sum(
        _replace_paragraph_xml(paragraph, target, replacement, split_only) for paragraph in root.iter(_PARAGRAPH_TAG)
    )
    if not count:
        return content, 0
    return etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True), count


def _replace_paragraph_xml(paragraph: etree._Element, target: str, replacement: str, split_only: bool) -> int:
    text_elements = [element for element in paragraph.iter(_TEXT_TAG)]
    full_text = "".join(element.text or "" for element in text_elements)
    count = 0
    position = full_text.find(target)
    while position >= 0:
        end = position + len(target)
        if not split_only or _spans_multiple_elements(text_elements, position, end):
            _replace_text_slice(text_elements, position, end, replacement)
            count += 1
        full_text = "".join(element.text or "" for element in text_elements)
        position = full_text.find(target, position + len(replacement))
    return count


def _spans_multiple_elements(text_elements: list, start: int, end: int) -> bool:
    touched = 0
    cursor = 0
    for element in text_elements:
        next_cursor = cursor + len(element.text or "")
        touched += int(next_cursor > start and cursor < end)
        cursor = next_cursor
    return touched > 1


def _replace_text_slice(text_elements: list, start: int, end: int, replacement: str) -> None:
    cursor = 0
    inserted = False
    for element in text_elements:
        text = element.text or ""
        next_cursor = cursor + len(text)
        if next_cursor <= start or cursor >= end:
            cursor = next_cursor
            continue
        element.text = _slice_text(text, cursor, start, end, replacement if not inserted else "")
        inserted = True
        cursor = next_cursor


def _slice_text(text: str, cursor: int, start: int, end: int, replacement: str) -> str:
    local_start = max(start - cursor, 0)
    local_end = min(end - cursor, len(text))
    return f"{text[:local_start]}{replacement}{text[local_end:]}"
