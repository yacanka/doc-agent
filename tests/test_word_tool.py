from pathlib import Path

from docx import Document

from app.tools.validation_tool import validate_replacements
from app.tools.word_tool import extract_text, replace_text, replace_text_preserve_style, safe_filename


def _docx(path: Path, text: str) -> None:
    document = Document()
    document.add_paragraph(text)
    document.save(path)


def _split_run_docx(path: Path) -> None:
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("Hello Al")
    paragraph.add_run("ice")
    document.save(path)


def test_safe_filename_removes_paths_and_symbols() -> None:
    assert safe_filename("../bad name!.docx") == "bad_name_.docx"


def test_extract_and_replace_docx_text(tmp_path: Path) -> None:
    source = tmp_path / "source.docx"
    output = tmp_path / "output.docx"
    _docx(source, "Hello Alice")

    count = replace_text(source, output, {"Alice": "Bob"})

    assert count == 1
    assert "Hello Bob" in extract_text(output)
    assert validate_replacements(output, {"Alice": "Bob"}, count)["valid"] is True


def test_replace_text_preserve_style_uses_timestamped_copy(tmp_path: Path) -> None:
    source = tmp_path / "source.docx"
    requested = tmp_path / "edited.docx"
    _docx(source, "Hello Alice")

    result = replace_text_preserve_style(source, "Alice", "Bob", requested)

    assert result.changed_count == 1
    assert result.output_path != requested
    assert result.output_path.name.endswith("_edited.docx")
    assert "Hello Bob" in extract_text(result.output_path)
    assert "Hello Alice" in extract_text(source)


def test_replace_text_preserve_style_handles_split_runs(tmp_path: Path) -> None:
    source = tmp_path / "split.docx"
    requested = tmp_path / "split-edited.docx"
    _split_run_docx(source)

    result = replace_text_preserve_style(source, "Alice", "Bob", requested)

    assert result.changed_count == 1
    assert "Hello Bob" in extract_text(result.output_path)
