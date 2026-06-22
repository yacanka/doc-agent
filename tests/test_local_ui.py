from pathlib import Path

import pytest


def test_ui_assets_are_localhost_relative() -> None:
    for asset_path in ["app/ui/index.html", "app/ui/app.js", "app/ui/styles.css"]:
        content = Path(asset_path).read_text(encoding="utf-8")
        assert "https://" not in content
        assert "http://" not in content


def test_ui_static_entrypoint_contains_required_panels() -> None:
    content = Path("app/ui/index.html").read_text(encoding="utf-8")

    expected_panels = ["Upload one document", "Chat", "Document info", "Developer stream", "Operation preview", "Download output"]
    for expected in expected_panels:
        assert expected in content


def test_ui_contains_streaming_and_loading_hooks() -> None:
    content = Path("app/ui/app.js").read_text(encoding="utf-8")

    for expected in ["questions/stream", "plan/stream", "readEventStream", "setLoading", "developerStream"]:
        assert expected in content


@pytest.fixture(name="client")
def fixture_client():
    pytest.importorskip("fastapi")
    pytest.importorskip("multipart")
    testclient = pytest.importorskip("fastapi.testclient")
    from app.main import app

    return testclient.TestClient(app)


def test_root_redirects_to_ui(client) -> None:
    response = client.get("/", follow_redirects=False)

    assert response.status_code in {307, 308}
    assert response.headers["location"] == "/ui/index.html"


def test_health_reports_offline_readiness(client) -> None:
    from app.main import config

    response = client.get("/health")

    assert response.status_code == 200
    assert response.json()["config"]["host"] == config.host
    assert "offline_ready" in response.json()


def test_download_output_uses_controlled_directory(client) -> None:
    from app.main import config

    output = config.outputs_dir / "ui-test-output.txt"
    output.write_text("done", encoding="utf-8")

    response = client.get("/outputs/ui-test-output.txt")

    assert response.status_code == 200
    assert response.text == "done"


def test_localhost_middleware_rejects_remote_clients() -> None:
    pytest.importorskip("fastapi")
    pytest.importorskip("multipart")
    testclient = pytest.importorskip("fastapi.testclient")
    from app.main import app

    with testclient.TestClient(app, client=("203.0.113.10", 1234)) as remote_client:
        response = remote_client.get("/health")

    assert response.status_code == 403


def test_question_stream_endpoint_emits_sse(client, monkeypatch, tmp_path) -> None:
    from app.main import store

    source = tmp_path / "sample.docx"
    from docx import Document
    document = Document()
    document.add_paragraph("Hello Alice")
    document.save(source)
    session = store.create_session("sample.docx", source)

    class FakeLlm:
        def stream_complete(self, prompt):
            yield "Hel"
            yield "lo"

    monkeypatch.setattr("app.main.llm", FakeLlm())

    response = client.post(f"/sessions/{session['id']}/questions/stream", json={"question": "Hi"})

    assert response.status_code == 200
    assert "event: token" in response.text
    assert 'data: {"text": "Hel"}' in response.text
    assert "event: done" in response.text


def test_plan_stream_endpoint_emits_tokens_and_replacements(client, monkeypatch, tmp_path) -> None:
    from app.main import store

    source = tmp_path / "sample.docx"
    from docx import Document
    document = Document()
    document.add_paragraph("Hello Alice")
    document.save(source)
    session = store.create_session("sample.docx", source)

    class FakeLlm:
        def stream_complete(self, prompt):
            yield 'Sure: {"replacements":{"Alice":"Bob"}}'

    monkeypatch.setattr("app.main.llm", FakeLlm())

    response = client.post(f"/sessions/{session['id']}/plan/stream", json={"instruction": "Replace Alice"})

    assert response.status_code == 200
    assert "event: token" in response.text
    assert "event: replacements" in response.text
    assert 'data: {"items": {"Alice": "Bob"}}' in response.text


def test_plan_endpoint_returns_422_for_invalid_model_json(client, monkeypatch, tmp_path) -> None:
    from app.main import store

    source = tmp_path / "sample.docx"
    from docx import Document
    document = Document()
    document.add_paragraph("Hello Alice")
    document.save(source)
    session = store.create_session("sample.docx", source)

    class FakePlanner:
        def plan(self, instruction, document_text):
            raise ValueError("Planner response must contain valid JSON")

    monkeypatch.setattr("app.main.planner", FakePlanner())

    response = client.post(f"/sessions/{session['id']}/plan", json={"instruction": "Replace Alice"})

    assert response.status_code == 422
    assert response.json()["detail"] == "Planner response must contain valid JSON"


def test_upload_accepts_xlsx_document(client, tmp_path) -> None:
    from openpyxl import Workbook

    source = tmp_path / "sample.xlsx"
    workbook = Workbook()
    workbook.active["A1"] = "Hello Alice"
    workbook.save(source)

    with source.open("rb") as handle:
        response = client.post(
            "/documents",
            files={"file": ("sample.xlsx", handle, "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")},
        )

    assert response.status_code == 200
    assert response.json()["filename"] == "sample.xlsx"
    assert "Hello Alice" in response.json()["text_preview"]
